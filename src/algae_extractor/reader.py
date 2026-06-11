import io
import re
import tempfile
import xml.etree.ElementTree as ET
from zipfile import ZipFile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _normalize_text_and_styles(chars: list[tuple[str, int]]) -> tuple[str, list[int]]:
    """
    Normalize whitespace like `normalize_whitespace`, but also return per-character
    style flags aligned to the normalized text.

    Style flags are bit-packed:
      - bit 1: italic
      - bit 2: bold
      - bit 0: neutral
    """
    out_chars: list[str] = []
    out_styles: list[int] = []
    prev_was_space = False

    for ch, style in chars:
        is_space = ch.isspace()
        if is_space:
            if not prev_was_space and out_chars:
                out_chars.append(" ")
                out_styles.append(0)  # neutral style for whitespace
            prev_was_space = True
            continue

        out_chars.append(ch)
        out_styles.append(style)
        prev_was_space = False

    # Trim trailing space to match `.strip()` semantics.
    if out_chars and out_chars[-1] == " ":
        out_chars.pop()
        out_styles.pop()

    return ("".join(out_chars), out_styles)


# Adobe "Symbol" font maps Latin code points to Greek letters / math glyphs
# (e.g. the byte for "m" renders as the micron µ, "p" renders as π). Word stores
# the underlying Latin character, so plain text extraction yields "m"/"p" unless
# we remap by font. Some Word builds instead store the glyph in the F0xx Private
# Use Area (0xF000 + code); we strip that offset before mapping.
_ADOBE_SYMBOL_TO_UNICODE: dict[int, str] = {
    # lowercase Greek (0x61-0x7A)
    0x61: "α", 0x62: "β", 0x63: "χ", 0x64: "δ", 0x65: "ε", 0x66: "φ",
    0x67: "γ", 0x68: "η", 0x69: "ι", 0x6A: "ϕ", 0x6B: "κ", 0x6C: "λ",
    0x6D: "μ", 0x6E: "ν", 0x6F: "ο", 0x70: "π", 0x71: "θ", 0x72: "ρ",
    0x73: "σ", 0x74: "τ", 0x75: "υ", 0x76: "ϖ", 0x77: "ω", 0x78: "ξ",
    0x79: "ψ", 0x7A: "ζ",
    # uppercase Greek (0x41-0x5A)
    0x41: "Α", 0x42: "Β", 0x43: "Χ", 0x44: "Δ", 0x45: "Ε", 0x46: "Φ",
    0x47: "Γ", 0x48: "Η", 0x49: "Ι", 0x4A: "ϑ", 0x4B: "Κ", 0x4C: "Λ",
    0x4D: "Μ", 0x4E: "Ν", 0x4F: "Ο", 0x50: "Π", 0x51: "Θ", 0x52: "Ρ",
    0x53: "Σ", 0x54: "Τ", 0x55: "Υ", 0x56: "ς", 0x57: "Ω", 0x58: "Ξ",
    0x59: "Ψ", 0x5A: "Ζ",
    # common math glyphs in Symbol encoding
    0xB1: "±", 0xB3: "≥", 0xA3: "≤", 0xB9: "≠", 0xBB: "↔", 0xB4: "×",
    0xB8: "÷", 0xB0: "°", 0xA5: "∞", 0xD7: "≅", 0xBD: "|",
}


def _remap_symbol_text(text: str, is_symbol_font: bool) -> str:
    """Convert Symbol-font / F0xx-PUA code points to their real Unicode glyphs."""
    if not text:
        return text
    out: list[str] = []
    for ch in text:
        cp = ord(ch)
        if 0xF000 <= cp <= 0xF0FF:
            # F0xx is "Symbol glyph at code (cp - 0xF000)": map Greek/math via the
            # table, else fall back to the plain ASCII glyph (space, digits, etc.)
            # rather than leaking the invisible PUA character through.
            base = cp - 0xF000
            out.append(_ADOBE_SYMBOL_TO_UNICODE.get(base, chr(base)))
        elif is_symbol_font:
            out.append(_ADOBE_SYMBOL_TO_UNICODE.get(cp, ch))
        else:
            out.append(ch)
    return "".join(out)


def _run_font_is_symbol(run) -> bool:
    name = getattr(run.font, "name", None)
    return bool(name) and name.lower() == "symbol"


_SUPERSCRIPT_MAP: dict[str, str] = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "-": "⁻", "+": "⁺", "(": "⁽", ")": "⁾", "n": "ⁿ", "i": "ⁱ",
}

_SUBSCRIPT_MAP: dict[str, str] = {
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "-": "₋", "+": "₊", "(": "₍", ")": "₎",
    "a": "ₐ", "e": "ₑ", "i": "ᵢ", "j": "ⱼ", "k": "ₖ",
    "l": "ₗ", "m": "ₘ", "n": "ₙ", "o": "ₒ", "p": "ₚ",
    "r": "ᵣ", "s": "ₛ", "t": "ₜ", "u": "ᵤ", "v": "ᵥ", "x": "ₓ",
}


def _apply_script_map(text: str, mapping: dict[str, str]) -> str:
    return "".join(mapping.get(c, c) for c in text)


def _paragraph_to_plain_and_styles(paragraph: Paragraph) -> tuple[str, list[int]] | None:
    chars: list[tuple[str, int]] = []
    for run in paragraph.runs:
        bold = bool(getattr(run, "bold", False))
        italic = bool(getattr(run, "italic", False))
        style = (1 if italic else 0) | (2 if bold else 0)
        if not run.text:
            continue
        run_text = _remap_symbol_text(run.text, _run_font_is_symbol(run))
        if getattr(run.font, "superscript", None):
            run_text = _apply_script_map(run_text, _SUPERSCRIPT_MAP)
        elif getattr(run.font, "subscript", None):
            run_text = _apply_script_map(run_text, _SUBSCRIPT_MAP)
        chars.extend([(c, style) for c in run_text])

    if not chars:
        return None

    plain, styles = _normalize_text_and_styles(chars)
    if not plain:
        return None
    return plain, styles


def paragraph_clean_text(paragraph: Paragraph) -> str:
    """Plain paragraph text with Symbol-font glyphs and super/subscripts normalized.

    Returns "" for empty paragraphs (unlike `_paragraph_to_plain_and_styles`, which
    returns None) so callers can join paragraphs line-by-line.
    """
    converted = _paragraph_to_plain_and_styles(paragraph)
    return converted[0] if converted else ""


def iter_docx_paragraphs(docx_path: str | Path):
    document = Document(str(docx_path))
    for paragraph in document.paragraphs:
        converted = _paragraph_to_plain_and_styles(paragraph)
        if not converted:
            continue
        text, char_styles = converted
        yield {
            "text": text,
            "char_styles": char_styles,
            "style": getattr(paragraph.style, "name", "") if paragraph.style else "",
        }


def _iter_document_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def table_to_row_texts(table: Table) -> list[list[str]]:
    """Plain text per cell, one row per Word table row (document order)."""
    rows_out: list[list[str]] = []
    for row in table.rows:
        cells_out: list[str] = []
        for cell in row.cells:
            parts: list[str] = []
            for paragraph in cell.paragraphs:
                converted = _paragraph_to_plain_and_styles(paragraph)
                if converted:
                    parts.append(converted[0])
            cells_out.append(normalize_whitespace(" ".join(parts)))
        rows_out.append(cells_out)
    return rows_out


_CHART_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
}

_SERIES_COLORS = [
    (31, 119, 180),
    (255, 127, 14),
    (44, 160, 44),
    (214, 39, 40),
    (148, 103, 189),
    (140, 86, 75),
]


def _safe_float(value: str | None) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _num_cache_values(parent: ET.Element | None) -> list[float]:
    if parent is None:
        return []
    out: list[float] = []
    for pt in parent.findall(".//c:numCache/c:pt", _CHART_NS):
        v = pt.find("c:v", _CHART_NS)
        f = _safe_float(v.text if v is not None else None)
        if f is not None:
            out.append(f)
    return out


def _chart_series_points(chart_root: ET.Element) -> list[list[tuple[float, float]]]:
    all_series: list[list[tuple[float, float]]] = []

    # Prefer scatter chart semantics if available.
    scatter = chart_root.find(".//c:scatterChart", _CHART_NS)
    if scatter is not None:
        for ser in scatter.findall("c:ser", _CHART_NS):
            x_values = _num_cache_values(ser.find("c:xVal", _CHART_NS))
            y_values = _num_cache_values(ser.find("c:yVal", _CHART_NS))
            points = list(zip(x_values, y_values))
            if points:
                all_series.append(points)
        if all_series:
            return all_series

    # Fallback for other chart families that use category/value pairs.
    for ser in chart_root.findall(".//c:ser", _CHART_NS):
        x_values = _num_cache_values(ser.find("c:cat", _CHART_NS))
        y_values = _num_cache_values(ser.find("c:val", _CHART_NS))
        points = list(zip(x_values, y_values))
        if points:
            all_series.append(points)

    return all_series


def _render_chart_to_png(chart_blob: bytes) -> bytes | None:
    try:
        root = ET.fromstring(chart_blob)
    except ET.ParseError:
        return None

    series = _chart_series_points(root)
    if not series:
        return None

    all_points = [pt for seq in series for pt in seq]
    if not all_points:
        return None

    x_vals = [p[0] for p in all_points]
    y_vals = [p[1] for p in all_points]
    x_min, x_max = min(x_vals), max(x_vals)
    y_min, y_max = min(y_vals), max(y_vals)
    if x_min == x_max:
        x_max = x_min + 1.0
    if y_min == y_max:
        y_max = y_min + 1.0

    width, height = 1400, 840
    margin_left, margin_top, margin_right, margin_bottom = 105, 40, 35, 85
    plot_w = width - margin_left - margin_right
    plot_h = height - margin_top - margin_bottom

    image = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(image)

    # Plot frame + subtle horizontal grid for readability.
    frame_color = (70, 70, 70)
    grid_color = (228, 232, 238)
    draw.rectangle(
        [margin_left, margin_top, margin_left + plot_w, margin_top + plot_h],
        outline=frame_color,
        width=2,
    )
    for i in range(1, 5):
        y = margin_top + int(plot_h * (i / 5))
        draw.line(
            [(margin_left, y), (margin_left + plot_w, y)],
            fill=grid_color,
            width=1,
        )

    def x_to_px(x: float) -> int:
        return margin_left + int((x - x_min) / (x_max - x_min) * plot_w)

    def y_to_px(y: float) -> int:
        return margin_top + int((1.0 - (y - y_min) / (y_max - y_min)) * plot_h)

    for idx, seq in enumerate(series):
        color = _SERIES_COLORS[idx % len(_SERIES_COLORS)]
        pts = [(x_to_px(x), y_to_px(y)) for x, y in seq]
        if len(pts) > 1:
            draw.line(pts, fill=color, width=2)
        for x, y in pts:
            draw.ellipse((x - 2, y - 2, x + 2, y + 2), fill=color)

    # Minimal axis labels: numeric min/max values.
    draw.text((margin_left, height - margin_bottom + 12), f"{x_min:.0f}", fill=frame_color)
    draw.text((margin_left + plot_w - 40, height - margin_bottom + 12), f"{x_max:.0f}", fill=frame_color)
    draw.text((12, margin_top + plot_h - 8), f"{y_min:.2g}", fill=frame_color)
    draw.text((12, margin_top - 8), f"{y_max:.2g}", fill=frame_color)

    out = io.BytesIO()
    image.save(out, format="PNG")
    return out.getvalue()


def _export_word_chart_images(docx_path: str | Path) -> list[bytes]:
    """
    Export chart objects from Word via COM by saving filtered HTML and reading
    chart image assets referenced as <img id="Chart N" ...>.
    Returns chart image blobs in HTML appearance order.
    """
    try:
        import pythoncom
        import win32com.client  # type: ignore[import-untyped]
    except Exception:
        return []

    path = str(Path(docx_path).resolve())
    chart_images: list[bytes] = []
    app = None
    doc = None
    tmpdir = tempfile.TemporaryDirectory()
    try:
        pythoncom.CoInitialize()
        app = win32com.client.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        doc = app.Documents.Open(path, ReadOnly=True, AddToRecentFiles=False)

        html_path = Path(tmpdir.name) / "word-export.html"
        # wdFormatFilteredHTML = 10
        doc.SaveAs2(str(html_path), FileFormat=10)

        assets_dir = html_path.with_name(f"{html_path.stem}_files")
        if html_path.exists() and assets_dir.exists():
            html = html_path.read_text(encoding="utf-8", errors="ignore")
            for tag in re.findall(r"<img\b[^>]*>", html, flags=re.IGNORECASE):
                id_m = re.search(r'\bid="([^"]+)"', tag, flags=re.IGNORECASE)
                src_m = re.search(r'\bsrc="([^"]+)"', tag, flags=re.IGNORECASE)
                if not id_m or not src_m:
                    continue
                chart_id = id_m.group(1).strip()
                if not re.match(r"^Chart\s+\d+$", chart_id, flags=re.IGNORECASE):
                    continue
                src_name = Path(src_m.group(1)).name
                img_path = assets_dir / src_name
                if not img_path.exists():
                    continue
                blob = img_path.read_bytes()
                if blob:
                    chart_images.append(blob)
    except Exception:
        return []
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if app is not None:
                app.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        tmpdir.cleanup()

    return chart_images


def _count_chart_nodes_in_docx(docx_path: str | Path) -> int:
    """Count <c:chart ...> nodes in word/document.xml."""
    try:
        with ZipFile(str(docx_path)) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        return len(re.findall(r"<c:chart\b", xml))
    except Exception:
        return 0


def _round_to_quarter_turn(angle: float) -> int | None:
    rounded = int(round(angle / 90.0) * 90)
    if abs(angle - rounded) > 0.75:
        return None
    return rounded % 360


def _apply_word_drawing_transform(
    *, blob: bytes, extension: str, drawing_el
) -> tuple[bytes, str]:
    """
    Apply Word drawing transforms (rotation/flip) so extracted images keep
    the orientation users see in the .docx.
    """
    xfrms = drawing_el.xpath(".//*[local-name()='xfrm']")
    transform = None
    for x in xfrms:
        if x.get("rot") is not None or x.get("flipH") is not None or x.get("flipV") is not None:
            transform = x
            break
    if transform is None:
        return blob, extension

    rot_attr = transform.get("rot")
    flip_h = str(transform.get("flipH", "")).lower() in {"1", "true"}
    flip_v = str(transform.get("flipV", "")).lower() in {"1", "true"}

    clockwise_deg = 0.0
    if rot_attr is not None:
        try:
            clockwise_deg = (int(rot_attr) / 60000.0) % 360.0
        except ValueError:
            clockwise_deg = 0.0

    quarter_turn = _round_to_quarter_turn(clockwise_deg)
    should_rotate = quarter_turn not in (None, 0)
    if not should_rotate and not flip_h and not flip_v:
        return blob, extension

    try:
        with Image.open(io.BytesIO(blob)) as im:
            out = im.copy()
            if flip_h:
                out = out.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
            if flip_v:
                out = out.transpose(Image.Transpose.FLIP_TOP_BOTTOM)
            if should_rotate and quarter_turn is not None:
                # OOXML rot is clockwise; PIL rotate is counter-clockwise.
                out = out.rotate(-quarter_turn, expand=True)

            target_ext = extension.lower()
            if target_ext in {".jpg", ".jpeg"}:
                fmt = "JPEG"
                if out.mode in ("RGBA", "LA", "P"):
                    out = out.convert("RGB")
            elif target_ext == ".gif":
                fmt = "GIF"
            else:
                target_ext = ".png"
                fmt = "PNG"
                if out.mode == "CMYK":
                    out = out.convert("RGB")

            buf = io.BytesIO()
            out.save(buf, format=fmt)
            return buf.getvalue(), target_ext
    except Exception:
        return blob, extension


def _yield_images_from_drawing_element(
    drawing_el,
    document: Document,
    *,
    word_chart_blobs: list[bytes] | None = None,
    chart_state: dict[str, int] | None = None,
):
    emitted_any = False

    for blip in drawing_el.xpath(".//a:blip"):
        relation_id = blip.get(qn("r:embed"))
        if not relation_id:
            continue
        image_part = document.part.related_parts.get(relation_id)
        if not image_part:
            continue
        extension = Path(getattr(image_part, "filename", "")).suffix.lower() or ".png"
        blob, extension = _apply_word_drawing_transform(
            blob=image_part.blob,
            extension=extension,
            drawing_el=drawing_el,
        )
        yield {
            "type": "image",
            "blob": blob,
            "extension": extension,
        }
        emitted_any = True

    # Some Word figures are embedded as chart objects (c:chart) without a:blip.
    if emitted_any:
        return
    for chart_node in drawing_el.xpath(".//c:chart"):
        relation_id = chart_node.get(qn("r:id"))
        if not relation_id:
            continue

        # Preferred path: Word-native chart rendering (style-preserving) exported via COM.
        if word_chart_blobs and chart_state is not None:
            used_word_blob = False
            if len(word_chart_blobs) == 1 and chart_state.get("total", 0) > 1:
                # If Word emitted only one chart image, it usually corresponds to
                # the main Figure 1 chart near the end of the record section.
                # Apply it to the final chart occurrence and keep others on fallback.
                if chart_state.get("seen", 0) == chart_state.get("total", 1) - 1:
                    yield {
                        "type": "image",
                        "blob": word_chart_blobs[0],
                        "extension": ".png",
                    }
                    used_word_blob = True
            else:
                idx = chart_state.get("blob_index", 0)
                if idx < len(word_chart_blobs):
                    yield {
                        "type": "image",
                        "blob": word_chart_blobs[idx],
                        "extension": ".png",
                    }
                    chart_state["blob_index"] = idx + 1
                    used_word_blob = True

            if used_word_blob:
                chart_state["seen"] = chart_state.get("seen", 0) + 1
                emitted_any = True
                continue

        chart_part = document.part.related_parts.get(relation_id)
        if not chart_part:
            if chart_state is not None:
                chart_state["seen"] = chart_state.get("seen", 0) + 1
            continue
        chart_blob = getattr(chart_part, "blob", None)
        if not chart_blob:
            if chart_state is not None:
                chart_state["seen"] = chart_state.get("seen", 0) + 1
            continue
        rendered = _render_chart_to_png(chart_blob)
        if not rendered:
            if chart_state is not None:
                chart_state["seen"] = chart_state.get("seen", 0) + 1
            continue
        yield {
            "type": "image",
            "blob": rendered,
            "extension": ".png",
        }
        if chart_state is not None:
            chart_state["seen"] = chart_state.get("seen", 0) + 1
        emitted_any = True



def iter_docx_content_blocks(docx_path: str | Path, *, use_word_renderer: bool = False):
    """
    Yield content in document order: paragraph text, tables, page breaks, and images
    interleave as in the WordprocessingML (e.g. figure then caption; page breaks
    between species).
    """
    document = Document(str(docx_path))
    word_chart_blobs = _export_word_chart_images(docx_path) if use_word_renderer else []
    total_chart_count = _count_chart_nodes_in_docx(docx_path)
    chart_state = {"blob_index": 0, "seen": 0, "total": total_chart_count}

    for block in _iter_document_blocks(document):
        if isinstance(block, Table):
            rows = table_to_row_texts(block)
            if any(any(c.strip() for c in r) for r in rows):
                yield {"type": "table", "rows": rows}
            continue

        if not hasattr(block, "runs"):
            continue

        paragraph = block
        buf_chars: list[tuple[str, int]] = []
        para_style = getattr(paragraph.style, "name", "") if paragraph.style else ""

        def take_paragraph_dict() -> dict[str, object] | None:
            if not buf_chars:
                return None
            plain, styles = _normalize_text_and_styles(buf_chars)
            buf_chars.clear()
            if not plain:
                return None
            return {
                "type": "paragraph",
                "text": plain,
                "char_styles": styles,
                "style": para_style,
            }

        for run in paragraph.runs:
            bold = bool(getattr(run, "bold", False))
            italic = bool(getattr(run, "italic", False))
            style_int = (1 if italic else 0) | (2 if bold else 0)
            is_symbol = _run_font_is_symbol(run)
            is_superscript = bool(getattr(run.font, "superscript", None))
            is_subscript = bool(getattr(run.font, "subscript", None))

            for el in run._element:
                tag = el.tag.split("}")[-1]

                if tag == "drawing" or tag == "pict":
                    sent = take_paragraph_dict()
                    if sent is not None:
                        yield sent
                    yield from _yield_images_from_drawing_element(
                        el,
                        document,
                        word_chart_blobs=word_chart_blobs,
                        chart_state=chart_state,
                    )
                elif tag == "AlternateContent":
                    # mc:AlternateContent wraps modern drawing markup (e.g. wpg:wgp groups).
                    # Use the mc:Choice (preferred) branch; fall back to mc:Fallback only if
                    # mc:Choice yields nothing.
                    emitted = False
                    for choice in el:
                        choice_local = choice.tag.split("}")[-1]
                        if choice_local == "Choice":
                            for inner in choice:
                                inner_tag = inner.tag.split("}")[-1]
                                if inner_tag in ("drawing", "pict"):
                                    sent = take_paragraph_dict()
                                    if sent is not None:
                                        yield sent
                                    yield from _yield_images_from_drawing_element(
                                        inner,
                                        document,
                                        word_chart_blobs=word_chart_blobs,
                                        chart_state=chart_state,
                                    )
                                    emitted = True
                            break
                    if not emitted:
                        for fallback in el:
                            fallback_local = fallback.tag.split("}")[-1]
                            if fallback_local == "Fallback":
                                for inner in fallback:
                                    inner_tag = inner.tag.split("}")[-1]
                                    if inner_tag in ("drawing", "pict"):
                                        sent = take_paragraph_dict()
                                        if sent is not None:
                                            yield sent
                                        yield from _yield_images_from_drawing_element(
                                            inner,
                                            document,
                                            word_chart_blobs=word_chart_blobs,
                                            chart_state=chart_state,
                                        )
                                break
                elif tag == "br":
                    br_type = el.get(qn("w:type"))
                    sent = take_paragraph_dict()
                    if sent is not None:
                        yield sent
                    if br_type == "page":
                        yield {"type": "page_break"}
                elif tag == "t":
                    if el.text:
                        run_text = _remap_symbol_text(el.text, is_symbol)
                        if is_superscript:
                            run_text = _apply_script_map(run_text, _SUPERSCRIPT_MAP)
                        elif is_subscript:
                            run_text = _apply_script_map(run_text, _SUBSCRIPT_MAP)
                        buf_chars.extend((ch, style_int) for ch in run_text)

        sent = take_paragraph_dict()
        if sent is not None:
            yield sent
